"""
generate_docx.py
================
Converts capstone_report.md to capstone_report.docx
Run: python3 reports/generate_docx.py
"""

from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))  # 1cm ≈ 567 twips
    trPr.append(trHeight)


def add_page_border(doc):
    """Add a subtle left border to code paragraphs via direct XML (done per paragraph)."""
    pass  # handled per paragraph


def style_doc(doc):
    """Apply global document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x2A)

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)


def add_heading(doc, text, level):
    colors = {
        1: RGBColor(0x0E, 0x3A, 0x6E),   # dark navy
        2: RGBColor(0x0E, 0x4D, 0x92),   # corporate blue
        3: RGBColor(0x00, 0x6B, 0x7F),   # teal
        4: RGBColor(0x33, 0x66, 0x99),   # steel blue
    }
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = colors.get(level, RGBColor(0x1A, 0x1A, 0x2A))
    if level == 1:
        p.runs[0].font.size = Pt(20)
        p.runs[0].font.bold = True
    elif level == 2:
        p.runs[0].font.size = Pt(16)
    elif level == 3:
        p.runs[0].font.size = Pt(14)
    elif level == 4:
        p.runs[0].font.size = Pt(12)
    return p


def add_paragraph(doc, text, bold=False, italic=False, color=None, size=11,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_code_block(doc, text):
    """Add a code block with monospace font and shaded background."""
    for line in text.split("\n"):
        p = doc.add_paragraph(style="No Spacing")
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x00, 0x40, 0x00)
        # Add shading via XML
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F0F4F0")
        pPr.append(shd)
    doc.add_paragraph()


def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header row."""
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, "0E4D92")
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = "EAF1FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            # Highlight bold markers
            if val.startswith("**") or "XGBoost" in val:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x0E, 0x3A, 0x6E)

    # Set column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def inline_format(p, text):
    """
    Add a run to paragraph `p`, handling **bold** and *italic* inline markers.
    Returns the paragraph.
    """
    # Split on **bold** and *italic* markers
    segments = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for seg in segments:
        if seg.startswith("**") and seg.endswith("**"):
            run = p.add_run(seg[2:-2])
            run.bold = True
        elif seg.startswith("*") and seg.endswith("*"):
            run = p.add_run(seg[1:-1])
            run.italic = True
        elif seg.startswith("`") and seg.endswith("`"):
            run = p.add_run(seg[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        elif seg:
            p.add_run(seg)
    return p


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def build_cover(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Capstone Report")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0E, 0x3A, 0x6E)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Network Traffic Anomaly Detection")
    run2.font.size = Pt(20)
    run2.font.color.rgb = RGBColor(0x0E, 0x4D, 0x92)

    doc.add_paragraph()
    meta = [
        ("Pillar", "5 — AI/ML Fundamentals Program"),
        ("Domain", "Cybersecurity — Network Intrusion Detection"),
        ("Dataset", "NSL-KDD (University of New Brunswick, 2009)"),
        ("Best Model", "XGBoost  |  F1 = 0.993  |  AUC-ROC = 0.999"),
        ("Date", "July 2025"),
    ]
    for k, v in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{k}:  ")
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0x0E, 0x3A, 0x6E)
        r2 = p.add_run(v)
        r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()


def build_toc(doc):
    add_heading(doc, "Table of Contents", 1)
    sections = [
        "1. Problem Understanding & Framing",
        "2. Data Collection & Understanding",
        "3. Data Preprocessing, EDA & Feature Engineering",
        "4. Model Implementation & Comparison",
        "5. Bias & Fairness Analysis",
        "6. Final Presentation Summary",
        "7. GitHub Repository & Deployment",
        "8. Conclusion",
        "9. Generative AI Usage",
    ]
    for s in sections:
        p = doc.add_paragraph(s, style="List Number")
        p.runs[0].font.color.rgb = RGBColor(0x0E, 0x4D, 0x92)
    doc.add_page_break()


def build_section1(doc):
    add_heading(doc, "1. Problem Understanding & Framing", 1)

    add_heading(doc, "1.1 Business Context", 2)
    p = doc.add_paragraph()
    inline_format(p,
        "Modern enterprise networks process millions of connection events daily. "
        "Undetected intrusions cost organisations an average **USD 4.45M per breach** (IBM, 2023). "
        "Traditional Intrusion Detection Systems (IDS) rely on static signatures and cannot detect "
        "novel (zero-day) attacks, generating excessive false positives that overwhelm Security "
        "Operations Centre (SOC) teams.")

    add_heading(doc, "Opportunity: An ML-based anomaly detection system can:", 4)
    for bullet in [
        "Detect previously unseen attack patterns",
        "Reduce false-positive rates by learning normal baselines",
        "Provide probabilistic risk scores for alert prioritisation",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    add_heading(doc, "1.2 Problem Statement", 2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(
        "Given a network connection record described by 41 features (protocol, byte volumes, "
        "error rates, connection statistics), determine whether the connection is normal or an "
        "attack, and identify its category (DoS, Probe, R2L, or U2R)."
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x0E, 0x3A, 0x6E)

    add_heading(doc, "1.3 Task Type", 2)
    add_table(doc,
              ["Dimension", "Selection"],
              [
                  ["Primary", "Binary Classification (Normal vs. Attack)"],
                  ["Secondary", "Multi-class (4 attack categories)"],
                  ["Unsupervised", "Anomaly Detection (Isolation Forest + Autoencoder)"],
              ],
              col_widths=[2.0, 4.5])

    add_heading(doc, "1.4 Success Metrics", 2)
    add_table(doc,
              ["Metric", "Target", "Justification"],
              [
                  ["F1-Score", "≥ 0.97", "Balances precision/recall; critical for imbalanced attack data"],
                  ["AUC-ROC", "≥ 0.99", "Ranking quality across all decision thresholds"],
                  ["False Positive Rate", "≤ 0.02", "SOC capacity constraint"],
                  ["Recall", "≥ 0.98", "Missing an attack is costlier than a false alarm"],
              ],
              col_widths=[2.0, 1.2, 3.8])

    p = doc.add_paragraph()
    inline_format(p,
        "**Business KPIs:** 80% reduction in undetected breaches, 40% reduction in alert fatigue, "
        "MTTD < 60 seconds.")


def build_section2(doc):
    add_heading(doc, "2. Data Collection & Understanding", 1)

    add_heading(doc, "2.1 Dataset", 2)
    p = doc.add_paragraph()
    inline_format(p, "**NSL-KDD** — University of New Brunswick (2009)")
    doc.add_paragraph("Source: https://www.unb.ca/cic/datasets/nsl.html")

    p = doc.add_paragraph()
    inline_format(p, "NSL-KDD improves upon the KDD Cup 1999 dataset by:")
    for b in [
        "Eliminating ~78% duplicate records (which distorted model performance)",
        "Providing balanced representation of rare attack types in the test set",
        "Enabling meaningful comparisons with published research",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    add_table(doc,
              ["Split", "Records", "Normal %", "Attack %"],
              [
                  ["Train", "125,973", "53.4%", "46.6%"],
                  ["Test", "22,544", "43.1%", "56.9%"],
              ],
              col_widths=[1.5, 1.5, 1.5, 1.5])

    add_heading(doc, "2.2 Feature Overview", 2)
    for b in [
        "41 features + label column",
        "38 numeric (int/float): byte counts, duration, error rates, connection statistics",
        "3 categorical: protocol_type (tcp/udp/icmp), service (70 values), flag (11 values)",
        "No missing values in NSL-KDD",
        "Attack categories: DoS (76%), Probe (19%), R2L (4%), U2R (0.3%)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        inline_format(p, b)

    add_heading(doc, "2.3 Key Features", 2)
    add_table(doc,
              ["Feature", "Security Meaning"],
              [
                  ["src_bytes / dst_bytes", "Volume of data — extremes indicate exfiltration or DoS"],
                  ["serror_rate", "Fraction of SYN-error connections — high = SYN flood"],
                  ["same_srv_rate", "Fraction of connections to same service — low = port scan"],
                  ["flag", "Connection status — S0 (no response) and REJ indicate attacks"],
                  ["logged_in", "Whether login succeeded — key predictor for R2L and U2R"],
              ],
              col_widths=[2.2, 4.8])


def build_section3(doc):
    add_heading(doc, "3. Data Preprocessing, EDA & Feature Engineering", 1)

    add_heading(doc, "3.1 Data Cleaning", 2)
    for b in [
        "Duplicates: 0 exact duplicates in NSL-KDD (by design)",
        "Missing values: None present (verified programmatically)",
        "Outliers: IQR-based Winsorisation applied using 3× IQR threshold — training data only to prevent leakage",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        inline_format(p, b)

    add_heading(doc, "3.2 Feature Engineering", 2)
    doc.add_paragraph("10 domain-derived features:")
    add_table(doc,
              ["Feature", "Formula / Source", "Security Rationale"],
              [
                  ["byte_ratio", "src_bytes / (dst_bytes + 1)", "Asymmetric traffic → C&C or exfiltration"],
                  ["total_bytes", "src_bytes + dst_bytes", "Total bandwidth consumed"],
                  ["log_src_bytes", "log1p(src_bytes)", "Normalises heavy-tailed distribution"],
                  ["log_dst_bytes", "log1p(dst_bytes)", "Normalises heavy-tailed distribution"],
                  ["error_rate_total", "serror_rate + rerror_rate", "Combined connection error indicator"],
                  ["srv_error_diff", "|serror_rate − srv_serror_rate|", "Service-level error anomaly"],
                  ["host_srv_ratio", "dst_host_srv_count / (dst_host_count+1)", "Service concentration at destination"],
                  ["is_long_connection", "duration > 30s", "Long connections indicate tunnelling"],
                  ["is_big_transfer", "total_bytes > 50KB", "Large transfers may indicate exfiltration"],
              ],
              col_widths=[1.9, 2.5, 2.6])

    add_heading(doc, "3.3 Key EDA Findings", 2)
    findings = [
        "serror_rate perfectly separates DoS/SYN attacks from normal traffic",
        "flag=S0 (connection attempted, no response) appears almost exclusively in attacks",
        "src_bytes=0, dst_bytes=0 is a strong attack indicator (probe/scan traffic)",
        "count > 200 in 2 seconds with high error rate is the signature of DDoS",
        "PCA scree plot: 95% variance retained in 23 components from ~55 total features",
    ]
    for i, f in enumerate(findings, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f)

    add_heading(doc, "3.4 Feature Selection", 2)
    add_table(doc,
              ["Method", "Features Selected", "Top Features"],
              [
                  ["Filter (Mutual Information)", "20",
                   "serror_rate, flag_SF, flag_S0, dst_host_serror_rate, same_srv_rate"],
                  ["Embedded (SelectFromModel)", "18",
                   "Gradient Boosting median importance threshold"],
                  ["Agreement (both methods)", "15",
                   "High-confidence feature set used for final models"],
              ],
              col_widths=[2.2, 1.3, 3.5])

    add_heading(doc, "3.5 Dimensionality Reduction", 2)
    for b in [
        "PCA: 55 features → 23 principal components (95% variance retained)",
        "t-SNE: Clear visual separation of Normal vs. Attack clusters in 2D embedding",
    ]:
        doc.add_paragraph(b, style="List Bullet")


def build_section4(doc):
    add_heading(doc, "4. Model Implementation & Comparison", 1)

    add_heading(doc, "4.1 Models Implemented", 2)
    add_table(doc,
              ["#", "Model", "Type", "Key Parameters"],
              [
                  ["1", "Logistic Regression", "Supervised (baseline)", "C=1.0, max_iter=1000"],
                  ["2", "Random Forest", "Supervised (ensemble)", "n_estimators=200, GridSearchCV"],
                  ["3", "XGBoost ★", "Supervised (boosting)", "n_estimators=300, max_depth=6, lr=0.1"],
                  ["4", "Isolation Forest", "Unsupervised", "n_estimators=200, contamination=0.47"],
                  ["5", "Autoencoder", "Deep Learning (unsupervised)", "64→32→16→32→64, 30 epochs"],
              ],
              col_widths=[0.4, 1.8, 2.0, 3.0])

    add_heading(doc, "4.2 Results", 2)
    add_table(doc,
              ["Model", "Accuracy", "Precision", "Recall", "F1", "AUC-ROC"],
              [
                  ["Logistic Regression", "0.921", "0.910", "0.930", "0.920", "0.967"],
                  ["Random Forest", "0.991", "0.991", "0.991", "0.991", "0.999"],
                  ["XGBoost ★ (best)", "0.993", "0.993", "0.993", "0.993", "0.999"],
                  ["Isolation Forest", "0.874", "0.850", "0.910", "0.880", "0.942"],
                  ["Autoencoder", "0.902", "0.880", "0.930", "0.900", "0.961"],
              ],
              col_widths=[2.0, 1.0, 1.0, 1.0, 0.8, 1.0])

    add_heading(doc, "4.3 Model Selection: XGBoost", 2)
    p = doc.add_paragraph()
    inline_format(p, "**XGBoost** is selected as the production model:")
    for b in [
        "Highest F1 (0.993) and AUC-ROC (0.999) on test set",
        "Robust to class imbalance via scale_pos_weight",
        "Natively handles mixed feature types",
        "Fast inference (<1ms per record)",
        "SHAP-compatible for full explainability",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    add_heading(doc, "4.4 Reproducibility", 2)
    add_table(doc,
              ["Artefact", "Location"],
              [
                  ["All models (joblib)", "models/*.pkl"],
                  ["Preprocessing pipeline", "models/preprocessor.pkl"],
                  ["Training configs & hyperparameters", "models/configs.yaml"],
                  ["Random seed", "42 throughout"],
                  ["Cross-validation", "StratifiedKFold k=3"],
              ],
              col_widths=[3.0, 4.0])


def build_section5(doc):
    add_heading(doc, "5. Bias & Fairness Analysis", 1)

    add_heading(doc, "5.1 Explainability — SHAP", 2)
    doc.add_paragraph("Top 10 Global SHAP Features (XGBoost):")
    features = [
        ("serror_rate", "Primary DoS/Probe indicator — high value strongly predicts 'attack'"),
        ("same_srv_rate", "Low value = port scan — distinguishes probe from normal"),
        ("flag_S0 / flag_SF", "Connection completion status — top one-hot encoded feature"),
        ("dst_host_serror_rate", "Host-level SYN error accumulation"),
        ("src_bytes", "Zero or extreme values = anomalous — key DoS indicator"),
        ("dst_host_srv_count", "Scan diversity"),
        ("logged_in", "Login success is a strong 'normal' indicator"),
        ("srv_serror_rate", "Service-level error rate"),
        ("dst_bytes", "Response data volume"),
        ("diff_srv_rate", "Service diversity in scan patterns"),
    ]
    for i, (feat, desc) in enumerate(features, 1):
        p = doc.add_paragraph(style="List Number")
        r1 = p.add_run(f"{feat}  ")
        r1.font.bold = True
        r1.font.name = "Courier New"
        r1.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        p.add_run(f"— {desc}")

    add_heading(doc, "5.2 LIME", 2)
    doc.add_paragraph(
        "Local LIME explanations for individual flagged connections confirm the SHAP global findings. "
        "For a detected SYN flood, the top contributing features are serror_rate=1.0, flag=S0, "
        "and count=511 — all classic SYN flood signatures."
    )

    add_heading(doc, "5.3 Known Limitations", 2)
    add_table(doc,
              ["Issue", "Severity", "Mitigation Applied"],
              [
                  ["Class Imbalance (U2R: 0.3%)", "High", "class_weight balancing; future: SMOTE"],
                  ["Dataset Age (1999 traffic)", "High", "Document; recommend CICIDS2017 retraining"],
                  ["Overfitting (train F1=0.999)", "Medium", "CV tuning, depth limits"],
                  ["Data Leakage (difficulty_level)", "Prevented", "Explicitly excluded from all pipelines"],
              ],
              col_widths=[2.5, 1.2, 3.3])

    add_heading(doc, "5.4 Fairness Audit — Traffic Subgroups", 2)
    doc.add_paragraph(
        "Since NSL-KDD contains no demographic attributes, fairness is evaluated across "
        "network traffic subgroups (protocol_type, service, flag)."
    )
    add_table(doc,
              ["Subgroup", "DPD", "EOD_FPR", "EOD_TPR", "FPR_Ratio", "Status"],
              [
                  ["protocol_type", "0.031", "0.007", "0.013", "1.19", "✅ PASS"],
                  ["service", "0.044", "0.018", "0.021", "1.22", "✅ PASS"],
                  ["flag", "0.052", "0.021", "0.028", "1.24", "✅ PASS"],
              ],
              col_widths=[1.8, 0.8, 0.9, 0.9, 1.0, 1.0])

    p = doc.add_paragraph()
    inline_format(p,
        "**All subgroups below concern thresholds** (DPD > 0.05, EOD > 0.10, Ratio > 1.25). "
        "Slight variation in ICMP/UDP reflects genuine traffic pattern differences, not model bias.")

    add_heading(doc, "5.5 Mitigation Strategies", 2)
    mitigations = [
        "Class imbalance: Apply SMOTE or class_weight='balanced' for U2R/R2L categories",
        "Data leakage: difficulty_level permanently excluded; preprocessor fit on training data only",
        "Overfitting: StratifiedKFold CV + depth constraints; monitor train-test F1 gap",
        "Distribution shift: Retrain on CICIDS-2017/2018; implement Evidently AI drift monitoring",
        "FPR disparity: If ratio > 1.25, use per-group classification thresholds as post-processing",
    ]
    for m in mitigations:
        p = doc.add_paragraph(style="List Bullet")
        inline_format(p, m)


def build_section6(doc):
    add_heading(doc, "6. Final Presentation Summary", 1)

    add_heading(doc, "Technical Presentation (12 slides)", 2)
    slides = [
        "Title & Agenda",
        "Problem Framing & Task Type",
        "Dataset Overview & Data Dictionary",
        "Preprocessing Pipeline",
        "EDA — Key Visualisations (distributions, correlation heatmap)",
        "Feature Engineering & Selection",
        "Model Architecture Overview",
        "Results Comparison Table + ROC Curves",
        "SHAP Global Importances",
        "Bias Audit Results",
        "Model Limitations & Mitigations",
        "Deployment Architecture + Conclusion",
    ]
    for i, s in enumerate(slides, 1):
        doc.add_paragraph(f"Slide {i}: {s}", style="List Number")

    add_heading(doc, "Business Presentation (10 slides)", 2)
    slides_b = [
        "Executive Summary — The Threat",
        "Our Solution: AI-Powered Network Defence",
        "How It Works (non-technical)",
        "Results: Key Numbers",
        "Cost-Benefit Analysis (ROI)",
        "Risk Assessment",
        "Implementation Roadmap",
        "Comparison vs. Traditional IDS",
        "Compliance & Ethical AI",
        "Next Steps & Recommendations",
    ]
    for i, s in enumerate(slides_b, 1):
        doc.add_paragraph(f"Slide {i}: {s}", style="List Number")


def build_section7(doc):
    add_heading(doc, "7. GitHub Repository & Deployment", 1)

    add_heading(doc, "7.1 Repository Structure", 2)
    add_code_block(doc,
"""network-anomaly-detection/
├── README.md                   # Project overview, setup, results
├── DEPLOYMENT.md               # Step-by-step deployment guide
├── requirements.txt            # All dependencies with version bounds
├── train_and_save.py           # One-shot training + artefact pipeline
├── src/
│   ├── app.py                  # FastAPI deployment endpoint
│   ├── ui.py                   # Streamlit interactive dashboard
│   ├── genai_explainer.py      # LLM-powered analyst explanation
│   ├── models.py               # Model training + evaluation
│   ├── preprocessor.py         # Sklearn pipeline
│   ├── data_loader.py          # NSL-KDD loader
│   ├── feature_engineering.py  # Domain features + PCA
│   ├── explainability.py       # SHAP + LIME + PDP
│   └── bias_audit.py           # Fairness audit
├── notebooks/                  # 01–06 reproducible notebooks
├── data/raw/                   # KDDTrain+.txt, KDDTest+.txt
├── models/                     # Saved models + configs
└── reports/                    # Plots, CSV reports, capstone_report.md""")

    add_heading(doc, "7.2 Local Deployment", 2)

    add_heading(doc, "FastAPI REST API", 3)
    add_code_block(doc,
"""# One-time: train the model
python train_and_save.py

# Start the API
uvicorn src.app:app --host 0.0.0.0 --port 8000

# Health check
curl -X GET http://localhost:8000/health -H "X-From: capstone-report"
# {"status": "ok", "model_loaded": true, "preprocessor_loaded": true}""")

    add_heading(doc, "Streamlit UI", 3)
    add_code_block(doc, "streamlit run src/ui.py   # Opens http://localhost:8501")
    doc.add_paragraph(
        "Features: preset traffic scenarios (Normal, Port Scan, DoS), full 41-feature form, "
        "real-time attack probability gauge, and the AI Analyst Explanation panel "
        "(requires OPENAI_API_KEY)."
    )

    add_heading(doc, "7.3 MLOps Practices", 2)
    add_table(doc,
              ["Practice", "Implementation"],
              [
                  ["Config management", "models/configs.yaml — hyperparameters + best_model identifier"],
                  ["Reproducibility", "Fixed seed 42; pinned requirements.txt; saved preprocessor"],
                  ["Retraining pipeline", "train_and_save.py — single command to retrain"],
                  ["Health monitoring", "GET /health — Prometheus / UptimeRobot compatible"],
                  ["Request traceability", "X-From header required; logged at INFO level"],
                  ["Input validation", "Pydantic schema with range constraints on all 41 features"],
                  ["Drift detection", "Evidently AI DataDriftPreset (recommended); retrain trigger: drift > 0.1"],
                  ["Secret management", "OPENAI_API_KEY via .env — never hardcoded"],
              ],
              col_widths=[2.2, 5.0])


def build_section8(doc):
    add_heading(doc, "8. Conclusion", 1)

    doc.add_paragraph(
        "This capstone project delivered a full end-to-end ML pipeline for network traffic "
        "anomaly detection."
    )

    outcomes = [
        ("Best model (XGBoost)", "F1=0.993, AUC-ROC=0.999 — exceeds all target metrics"),
        ("5 models implemented", "Supervised and unsupervised approaches compared"),
        ("Explainability", "SHAP confirms model decisions align with security domain knowledge"),
        ("Bias audit", "No significant fairness concerns across traffic subgroups"),
        ("FastAPI deployment", "Real-time inference at <1ms latency"),
    ]
    for title, detail in outcomes:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(f"{title}: ")
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0x0E, 0x3A, 0x6E)
        p.add_run(detail)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Business Impact")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0E, 0x3A, 0x6E)
    doc.add_paragraph(
        "At a 0.7% FPR, an enterprise processing 1 million connections per day would receive "
        "~7,000 false alerts vs. ~50,000 from a typical signature IDS — a 6× reduction in alert "
        "fatigue, enabling SOC teams to focus on genuine threats."
    )


def build_section9(doc):
    add_heading(doc, "9. Generative AI Usage", 1)

    add_table(doc,
              ["Tool", "Usage", "Verification"],
              [
                  ["GitHub Copilot",
                   "Boilerplate code (sklearn pipelines, FastAPI schema), docstring suggestions",
                   "All code reviewed and tested; logic modified to match project requirements"],
                  ["GPT-4",
                   "Generated initial EDA summary text and data dictionary descriptions",
                   "Content verified against actual dataset statistics; all figures re-run from code"],
                  ["OpenAI GPT-4o-mini",
                   "Real-time SOC analyst explanations (built into Streamlit UI)",
                   "Prompt-engineered to use validated model output only; no hallucinated statistics"],
              ],
              col_widths=[1.8, 2.8, 2.6])

    add_heading(doc, "9.1 GenAI-Enhanced Feature: AI Analyst Explainer", 2)
    doc.add_paragraph(
        "src/genai_explainer.py implements a live LLM-powered explanation layer wired into the "
        "Streamlit UI. After each prediction, the user can expand the 'AI Analyst Explanation' "
        "panel to receive a plain-English security briefing generated by GPT-4o-mini."
    )

    add_heading(doc, "Example Output — DoS Attack (Neptune):", 3)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(
        '"This connection shows all hallmarks of a Neptune SYN flood: 511 connection attempts '
        'to the same HTTP service with a 100% SYN error rate and zero bytes transferred in either '
        'direction, indicating no TCP handshake was ever completed. The SOC analyst should '
        'immediately isolate the source IP, escalate to Tier 2, and apply a temporary block rule '
        'on the firewall."'
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x0E, 0x3A, 0x6E)

    add_heading(doc, "9.2 Responsible AI Use", 2)
    principles = [
        "No AI-generated code was used without understanding and verification",
        "All statistical claims were validated against actual notebook output",
        "AI suggestions were treated as starting points, not final answers",
        "Security-critical code (app.py input validation) was written manually and reviewed for OWASP issues",
        "The LLM explanation prompt operates on model-validated data only — cannot fabricate statistics",
    ]
    for p_text in principles:
        doc.add_paragraph(p_text, style="List Bullet")

    add_heading(doc, "What was NOT generated by AI:", 4)
    not_ai = [
        "Model architecture decisions",
        "Feature engineering domain logic (security expertise applied)",
        "Bias audit framework design",
        "Business KPI calculations",
    ]
    for item in not_ai:
        doc.add_paragraph(item, style="List Bullet")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()
    style_doc(doc)

    build_cover(doc)
    build_toc(doc)
    build_section1(doc)
    doc.add_page_break()
    build_section2(doc)
    doc.add_page_break()
    build_section3(doc)
    doc.add_page_break()
    build_section4(doc)
    doc.add_page_break()
    build_section5(doc)
    doc.add_page_break()
    build_section6(doc)
    doc.add_page_break()
    build_section7(doc)
    doc.add_page_break()
    build_section8(doc)
    doc.add_page_break()
    build_section9(doc)

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Report prepared in accordance with Pillar 5 Capstone instructions.\n"
        "All code is reproducible — run notebooks 01–06 in sequence after pip install -r requirements.txt"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    out = Path(__file__).parent / "capstone_report.docx"
    doc.save(str(out))
    print(f"✅  Saved: {out}")


if __name__ == "__main__":
    main()
